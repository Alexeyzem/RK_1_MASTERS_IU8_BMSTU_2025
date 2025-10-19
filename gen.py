from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def create_report_docx(filename="Отчет.docx"):
    document = Document()

    # Установка стиля для заголовка
    document.add_heading('Отчет об анализе коммерческой деятельности', 0)
    document.add_paragraph('Отдел: Отдел продаж (ID: 17)').bold = True

    document.add_heading('1. Анализ продаж и финансовые результаты', level=1)

    table_data_1 = [
        ['Показатель', 'Значение', 'Вывод'],
        ['Общая прибыль отдела (текущая)', '2,982,786 руб.',
         'Общая прибыль от всех проектов, в которых участвовал отдел.'],
        ['Средний ROI', '9.93%', 'Текущий возврат инвестиций по проектам отдела.'],
        ['Менеджеры с наибольшим числом завершенных проектов', 'Не удалось определить',
         'Проблема данных: Отсутствует связь "проект → менеджер".']
    ]

    table_1 = document.add_table(rows=len(table_data_1), cols=len(table_data_1[0]), style='Table Grid')
    for i, row in enumerate(table_data_1):
        for j, cell_text in enumerate(row):
            table_1.cell(i, j).text = cell_text
            if i == 0:
                table_1.cell(i, j).paragraphs[0].runs[0].bold = True

    document.add_heading('2. Персональная эффективность', level=1)

    table_data_2 = [
        ['Показатель', 'Значение', 'Вывод'],
        ['Revenue per Employee', '114,722.53 руб.', 'Показатель выручки на одного сотрудника отдела.'],
        ['Корреляция (Зарплата vs. Performance Score)', '0.25', 'Слабая положительная корреляция (25% связи).']
    ]

    table_2 = document.add_table(rows=len(table_data_2), cols=len(table_data_2[0]), style='Table Grid')
    for i, row in enumerate(table_data_2):
        for j, cell_text in enumerate(row):
            table_2.cell(i, j).text = cell_text
            if i == 0:
                table_2.cell(i, j).paragraphs[0].runs[0].bold = True

    document.add_heading('3. Языковая аналитика', level=1)

    language_needs = [
        ['Михайлов Виталий Николаевич', 'Русский'],
        ['Павлова Лидия Вадимовна', 'Русский и английский'],
        ['Виноградова Диана Вячеславовна', 'Английский'],
        ['Виноградов Геннадий Валерьевич', 'Русский и английский'],
        ['Петрова Галина Евгеньевна', 'Английский'],
        ['Федорова Анастасия Павловна', 'Английский'],
        ['Богданова Яна Владимировна', 'Русский'],
        ['Лебедев Антон Олегович', 'Русский и английский'],
        ['Морозова Варвара Михайловна', 'Русский и английский'],
        ['Морозова Светлана Вячеславовна', 'Русский'],
        ['Воробьев Никита Павлович', 'Русский'],
        ['Волков Вадим Владимирович', 'Русский'],
        ['Соловьева Юлия Максимовна', 'Английский'],
        ['Морозов Тимур Артемович', 'Английский'],
        ['Новикова Варвара Павловна', 'Русский и английский'],
        ['Новиков Ярослав Евгеньевич', 'Английский'],
        ['Беляева Татьяна Борисовна', 'Русский и английский'],
        ['Кузнецова Наталья Борисовна', 'Английский'],
        ['Иванов Игорь Кириллович', 'Русский'],
        ['Михайлова Ирина Ивановна', 'Английский'],
        ['Козлова Светлана Максимовна', 'Русский и английский'],
        ['Козлов Валерий Геннадьевич', 'Английский']
    ]

    document.add_heading('Потребность в дополнительном языковом обучении', level=2)

    # Таблица с сотрудниками и языками
    table_3 = document.add_table(rows=len(language_needs) + 1, cols=2, style='Table Grid')
    table_3.cell(0, 0).text = 'Сотрудник'
    table_3.cell(0, 1).text = 'Требуемые языки'
    for i in range(2):
        table_3.cell(0, i).paragraphs[0].runs[0].bold = True

    for i, (name, lang) in enumerate(language_needs):
        table_3.cell(i + 1, 0).text = name
        table_3.cell(i + 1, 1).text = lang

    document.add_heading('Распределение навыков', level=2)
    document.add_paragraph('Русский: 12, Китайский: 11, Английский: 11, Немецкий: 10, Французский: 10.')

    document.add_heading('Сотрудники со знанием английского и немецкого одновременно', level=2)
    for name in ['Семенов Борис Валерьевич', 'Морозова Светлана Вячеславовна', 'Воробьев Никита Павлович',
                 'Волков Вадим Владимирович', 'Иванов Игорь Кириллович', 'Васильева Валентина Николаевна']:
        document.add_paragraph(name, style='List Bullet')

    document.add_heading('4. Клиентская аналитика', level=1)
    document.add_heading('Проекты с высоким риском и высокой прибыльностью', level=2)
    document.add_paragraph('Проект Уничтожение 9107: Низкий приоритет. Прибыль — 568 234 руб.')
    document.add_paragraph('Проект Подробность 4877: Критичный приоритет. Прибыль — 1 475 663 руб.')

    document.add_heading('5. Стратегические рекомендации', level=1)
    document.add_paragraph(f'Потенциальный эффект от увеличения ROI на 5%: Увеличение прибыли на 511,441 руб.')

    document.add_heading('Целевые показатели на следующий квартал', level=2)
    for goal in ['Увеличение прибыли на 500,000 руб.', 'Увеличение ROI на 4% минимум',
                 'Записать на языковые курсы необходимых сотрудников']:
        document.add_paragraph(goal, style='List Number')

    document.save(filename)
    print(f"Отчет успешно сохранен в файл '{filename}'")


# Запуск функции
create_report_docx()
